
import re
import pandas as pd
from docx import Document
import streamlit as st
from io import BytesIO
import os
from datetime import date

MASTER_FILE = "MPH_Clinic_Master.xlsx"

def infer_clinic_type(plan):
    plan_lower = plan.lower()
    if "discharge" in plan_lower:
        return "Discharge"
    elif "review" in plan_lower or "follow-up" in plan_lower:
        return "Follow-Up"
    elif "consent" in plan_lower or "new patient" in plan_lower:
        return "New"
    elif "post rt" in plan_lower:
        return "Post-RT Review"
    return "Other"

def extract_rt_dose_fractions(rt_text):
    match = re.search(r'(\d+\.?\d*)Gy/(\d+)', rt_text)
    return match.groups() if match else ("", "")

def extract_boost(rt_text):
    match = re.search(r'\+\s*(\d+\.?\d*)Gy/(\d+)', rt_text)
    return match.groups() if match else ("", "")

def extract_hormone_therapy(plan_text):
    plan_lower = plan_text.lower()
    if "tamoxifen" in plan_lower:
        return "Tamoxifen"
    elif "letrozole" in plan_lower:
        return "Letrozole"
    return "None"

def extract_bone_support(plan_text):
    lower = plan_text.lower()
    if any(term in lower for term in ["zoledronic", "zometa", "alendronic", "adcal", "vitamin d"]):
        return "Yes"
    return "No"

def extract_endocrine_plan(plan_text):
    plan_lower = plan_text.lower()
    if "stop" in plan_lower and "tamoxifen" in plan_lower:
        return "Stopped Tamoxifen"
    elif "start" in plan_lower and "letrozole" in plan_lower:
        return "Started Letrozole"
    elif "switch" in plan_lower:
        return "Switched Endocrine"
    return ""

def parse_clinic_doc(doc, clinic_date, selected_clinic_type):
    document = Document(doc)
    patient_start_pattern = re.compile(r'^[A-Za-z\-\.\' ]+ \d{5,7}(?: [A-Za-z]+)?$')
    patient_indices = [i for i, p in enumerate(document.paragraphs) if patient_start_pattern.match(p.text.strip())]
    patient_indices.append(len(document.paragraphs))

    all_data = []

    for start, end in zip(patient_indices, patient_indices[1:]):
        section = [p.text.strip() for p in document.paragraphs[start:end] if p.text.strip()]
        plan_text = ""
        rt_text = ""
        data = {
            "Clinic Date": clinic_date,
            "Clinic Type": selected_clinic_type,
            "Patient Name": "",
            "Age": "",
            "Diagnosis": "",
            "Surgery Date": "",
            "Tumour Size (mm)": "",
            "Tumour Grade": "",
            "ER/PR/HER2": "",
            "NPI": "",
            "RT Course & Dose": "",
            "RT Dose (Gy)": "",
            "RT Fractions": "",
            "Boost Dose (Gy)": "",
            "Boost Fractions": "",
            "Endocrine Therapy": "",
            "Endocrine Plan": "",
            "Bone Health Support": "",
            "Final RT Date": "",
            "Skin Side Effects": "",
            "Pain Side Effects": "",
            "Fatigue Side Effects": "",
            "Hormone Side Effects": "",
            "Plan": ""
        }

        name_line = section[0]
        name_parts = name_line.split()
        id_index = next((i for i, x in enumerate(name_parts) if x.isdigit()), None)
        if id_index:
            data["Patient Name"] = " ".join(name_parts[:id_index])

        for line in section:
            if "years old" in line:
                if m := re.search(r'(\d+)\s+years old', line): data["Age"] = m.group(1)
                if m := re.search(r'\((\d{2}/\d{2}/\d{4})', line): data["Surgery Date"] = m.group(1)
                if m := re.search(r'(\d+(\.\d+)?)\s*mm,', line): data["Tumour Size (mm)"] = m.group(1)
                if m := re.search(r'\bG([1-3])\b', line): data["Tumour Grade"] = m.group(1)
                if "IDC" in line: data["Diagnosis"] = "IDC"
                elif "DCIS" in line: data["Diagnosis"] = "DCIS"
                if m := re.search(r'(ER [0-9]+.*?)(PR [0-9]+)?(,? HER2[^,]*)?', line):
                    er = m.group(1) or ""
                    pr = m.group(2) or ""
                    her2 = m.group(3) or ""
                    data["ER/PR/HER2"] = f"{er} {pr} {her2}".strip()
                if m := re.search(r'(NPI|VNPI):\s*([\d\.]+)', line): data["NPI"] = m.group(2)

        for line in section:
            if line.startswith("Final radiotherapy"):
                parts = re.split(r'[\u2013\u2014\-]', line)
                if len(parts) >= 2:
                    data["Final RT Date"] = parts[0].replace("Final radiotherapy", "").strip()
                    rt_text = parts[1].strip()
                    data["RT Course & Dose"] = rt_text
                    dose, fracs = extract_rt_dose_fractions(rt_text)
                    data["RT Dose (Gy)"] = dose
                    data["RT Fractions"] = fracs
                    boost_dose, boost_frac = extract_boost(rt_text)
                    data["Boost Dose (Gy)"] = boost_dose
                    data["Boost Fractions"] = boost_frac

        for line in section:
            if line.startswith("Plan:"):
                plan_text = line[5:].strip()
                data["Plan"] = plan_text

        if selected_clinic_type == "Auto-detect":
            data["Clinic Type"] = infer_clinic_type(plan_text)

        data["Endocrine Therapy"] = extract_hormone_therapy(plan_text)
        data["Bone Health Support"] = extract_bone_support(plan_text)
        data["Endocrine Plan"] = extract_endocrine_plan(plan_text)

        for line in section:
            if line.startswith("Skin: "): data["Skin Side Effects"] = line[6:].strip().split("/")[0]
            if line.startswith("Pain: "): data["Pain Side Effects"] = line[6:].strip().split("/")[0]
            if line.startswith("Fatigue: "): data["Fatigue Side Effects"] = line[9:].strip().split("/")[0]
            if line.lower().startswith("on tamoxifen") or line.lower().startswith("on letrozole"):
                data["Hormone Side Effects"] = "None" if "no SE" in line else line.split("which")[-1].strip()

        all_data.append(data)

    return pd.DataFrame(all_data)

# Streamlit UI
st.title("MPH Radiotherapy Clinic Audit Uploader")

if st.button("🗑️ Clear all audit data (reset file)"):
    if os.path.exists(MASTER_FILE):
        os.remove(MASTER_FILE)
        st.success("Audit file deleted successfully. You can start fresh.")
    else:
        st.info("No audit file found to delete.")

clinic_date = st.date_input("Select clinic date", value=date.today())
clinic_type = st.selectbox("Select clinic type", options=["Auto-detect", "New", "Follow-Up", "Post-RT Review", "Discharge", "Other"])
uploaded_file = st.file_uploader("Upload your .docx clinic file", type=["docx"])

if uploaded_file is not None:
    new_data = parse_clinic_doc(uploaded_file, clinic_date, clinic_type)
    st.success("File parsed successfully. Preview below:")
    st.dataframe(new_data)

    if os.path.exists(MASTER_FILE):
        existing_data = pd.read_excel(MASTER_FILE)
        combined_data = pd.concat([existing_data, new_data], ignore_index=True)
        combined_data.drop_duplicates(subset=["Patient Name", "Surgery Date", "Clinic Date"], keep="first", inplace=True)
    else:
        combined_data = new_data

    combined_data.to_excel(MASTER_FILE, index=False)

    to_download = BytesIO()
    combined_data.to_excel(to_download, index=False, engine='openpyxl')
    to_download.seek(0)
    st.download_button(
        label="Download Updated Master Audit File",
        data=to_download,
        file_name="MPH_Clinic_Master.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    st.info(f"Total records: {len(combined_data)} | Newly added: {len(new_data)}")
